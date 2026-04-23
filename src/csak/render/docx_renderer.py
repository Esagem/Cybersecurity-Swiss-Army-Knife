"""Docx renderer — walks the report context and emits a Word document
via python-docx.

First-pass priority is correct structure over typography: sensible
headings, readable defaults, no broken tables. Polish (cover page,
tight spacing, matched fonts) is a deliberate second pass once the
structure stabilizes.
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt

from csak.query.context import FindingView, MethodologyScan, ReportContext, TicketGroup


def write_report(ctx: ReportContext, out_path: Path) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _apply_default_styles(doc)

    if ctx.kind == "internal-review":
        _write_internal_review(doc, ctx)
    elif ctx.kind == "fit-bundle":
        _write_fit_bundle(doc, ctx)
    else:
        raise ValueError(f"unknown report kind: {ctx.kind}")

    doc.save(out_path)
    return out_path


def write_ticket(ticket: TicketGroup, out_path: Path) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _apply_default_styles(doc)
    _write_ticket(doc, ticket)
    doc.save(out_path)
    return out_path


def _apply_default_styles(doc: Document) -> None:
    # First pass: only set a readable body font size on the normal
    # style. python-docx preserves the template's heading styles.
    style = doc.styles["Normal"]
    style.font.size = Pt(11)


def _write_internal_review(doc: Document, ctx: ReportContext) -> None:
    doc.add_heading(f"Internal Review — {ctx.org.name}", level=1)

    meta = doc.add_paragraph()
    meta.add_run("Period: ").bold = True
    meta.add_run(
        f"{ctx.period.label}  "
        f"({_isodate(ctx.period.start)} → {_isodate(ctx.period.end)})"
    )
    meta2 = doc.add_paragraph()
    meta2.add_run("Generated: ").bold = True
    meta2.add_run(_isodatetime(ctx.generated_at))
    meta3 = doc.add_paragraph()
    meta3.add_run("CSAK: ").bold = True
    meta3.add_run(ctx.csak_version)

    # Summary.
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(f"Findings in window: {len(ctx.findings)}")
    doc.add_paragraph(f"Targets touched: {len(ctx.findings_by_target)}")
    doc.add_paragraph(f"Contributing scans: {len(ctx.methodology)}")

    _write_severity_table(doc, ctx)

    # Methodology.
    doc.add_heading("Methodology", level=2)
    if ctx.methodology:
        _write_methodology_table(doc, ctx.methodology)
        for m in ctx.methodology:
            if m.timestamp_disclaimer:
                p = doc.add_paragraph()
                r = p.add_run(f"Note on {m.scan.label}: ")
                r.bold = True
                p.add_run(m.timestamp_disclaimer)
    else:
        doc.add_paragraph("(no scans contributed to this period)")

    # Findings.
    doc.add_heading("Findings", level=2)
    if not ctx.findings:
        doc.add_paragraph("(no findings in this period)")
        return

    for severity, group in ctx.findings_by_severity.items():
        if not group:
            continue
        doc.add_heading(f"{severity.capitalize()} ({len(group)})", level=3)
        for v in group:
            _write_finding_card(doc, v)


def _write_severity_table(doc: Document, ctx: ReportContext) -> None:
    doc.add_heading("Severity breakdown", level=3)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Severity"
    hdr[1].text = "Count"
    for sev, items in ctx.findings_by_severity.items():
        if not items:
            continue
        row = table.add_row().cells
        row[0].text = sev
        row[1].text = str(len(items))


def _write_methodology_table(doc: Document, methodology: list[MethodologyScan]) -> None:
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Tool"
    hdr[1].text = "Label"
    hdr[2].text = "Started"
    hdr[3].text = "Completed"
    hdr[4].text = "Timestamp"
    hdr[5].text = "Findings"
    for m in methodology:
        row = table.add_row().cells
        row[0].text = m.scan.source_tool
        row[1].text = m.scan.label
        row[2].text = _isodatetime(m.scan.scan_started_at)
        row[3].text = _isodatetime(m.scan.scan_completed_at)
        row[4].text = m.scan.timestamp_source
        row[5].text = str(m.finding_count)


def _write_finding_card(doc: Document, v: FindingView) -> None:
    doc.add_heading(v.finding.title, level=4)
    f = v.finding
    labelled(doc, "Target", f"{v.target.name}")
    labelled(doc, "Tool", f.source_tool)
    priority_expr = (
        f"{f.priority:.3f}  "
        f"({f.severity_weight:.2f} × {f.confidence_weight:.2f} × "
        f"{v.target.target_weight:.2f} × {f.probability_real:.2f})"
    )
    labelled(doc, "Priority", priority_expr)
    labelled(doc, "Confidence", f.confidence)
    labelled(doc, "Status", f.status)
    labelled(doc, "First seen", _isodatetime(f.first_seen))
    labelled(doc, "Last seen", _isodatetime(f.last_seen))
    if v.scans:
        doc.add_paragraph("Observed in scans:", style="List Bullet").runs[0].bold = True
        for s in v.scans:
            doc.add_paragraph(f"{s.label} ({s.source_tool})", style="List Bullet 2")


def _write_fit_bundle(doc: Document, ctx: ReportContext) -> None:
    doc.add_heading(f"Fix-it Ticket Bundle — {ctx.org.name}", level=1)
    meta = doc.add_paragraph()
    meta.add_run("Period: ").bold = True
    meta.add_run(f"{ctx.period.label}")
    meta2 = doc.add_paragraph()
    meta2.add_run("Generated: ").bold = True
    meta2.add_run(_isodatetime(ctx.generated_at))
    doc.add_paragraph(
        f"This bundle contains {len(ctx.tickets)} ticket(s). Findings "
        f"affecting multiple assets are collapsed into one ticket with "
        f"all assets listed."
    )
    for t in ctx.tickets:
        _write_ticket(doc, t)
        doc.add_paragraph()  # separator


def _write_ticket(doc: Document, t: TicketGroup) -> None:
    doc.add_heading(f"{t.ticket_id} — {t.title}", level=2)
    labelled(doc, "Severity", t.severity or "needs analyst review")
    doc.add_paragraph("Affected assets:").runs[0].bold = True
    for asset in t.affected_assets:
        doc.add_paragraph(asset, style="List Bullet")

    doc.add_heading("Impact", level=3)
    doc.add_paragraph(t.impact)
    doc.add_heading("Remediation", level=3)
    doc.add_paragraph(t.remediation)
    doc.add_heading("Validation", level=3)
    doc.add_paragraph(t.validation)


def labelled(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.bold = True
    p.add_run(value)


def _isodate(d: datetime) -> str:
    return d.date().isoformat()


def _isodatetime(d: datetime) -> str:
    return d.strftime("%Y-%m-%d %H:%M:%S UTC")
