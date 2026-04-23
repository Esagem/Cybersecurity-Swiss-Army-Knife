from pathlib import Path

from docx import Document

from csak.ingest.pipeline import ingest_path
from csak.query.context import build_context
from csak.query.finders import parse_period
from csak.render import docx_renderer as dx
from csak.storage import repository as repo


NUCLEI = """\
{"template-id":"cve-x","info":{"name":"Bad lib","severity":"high"},"matched-at":"https://api.acme.com","host":"api.acme.com","timestamp":"2026-04-20T10:00:00Z"}
{"template-id":"banner","info":{"name":"Banner","severity":"info"},"matched-at":"https://acme.com","host":"acme.com","timestamp":"2026-04-20T10:02:00Z"}
"""


def _seed(db, tmp_path, artifacts_dir):
    org = repo.create_org(db, name="Acme Corp", slug="acme")
    p = tmp_path / "n.jsonl"
    p.write_text(NUCLEI, encoding="utf-8")
    ingest_path(
        db, org_id=org.id, source_tool="nuclei",
        path=p, artifacts_root=artifacts_dir,
    )
    return org


def test_internal_review_docx_has_expected_headings(
    db, tmp_path: Path, artifacts_dir: Path, reports_dir: Path
) -> None:
    org = _seed(db, tmp_path, artifacts_dir)
    ctx = build_context(db, org=org, period=parse_period("2026-04"), kind="internal-review")
    out = dx.write_report(ctx, reports_dir / "r.docx")
    assert out.exists()

    doc = Document(str(out))
    headings = [
        p.text for p in doc.paragraphs if p.style.name.startswith("Heading")
    ]
    assert any("Internal Review" in h for h in headings)
    assert any("Summary" in h for h in headings)
    assert any("Methodology" in h for h in headings)
    assert any("Findings" in h for h in headings)


def test_internal_review_docx_includes_severity_table(
    db, tmp_path: Path, artifacts_dir: Path, reports_dir: Path
) -> None:
    org = _seed(db, tmp_path, artifacts_dir)
    ctx = build_context(db, org=org, period=parse_period("2026-04"), kind="internal-review")
    out = dx.write_report(ctx, reports_dir / "r.docx")

    doc = Document(str(out))
    # Two tables expected: severity breakdown + methodology.
    assert len(doc.tables) >= 2
    sev_cells = [row.cells[0].text for row in doc.tables[0].rows]
    assert "Severity" in sev_cells
    assert "high" in sev_cells
    assert "info" in sev_cells


def test_fit_bundle_docx_includes_tickets(
    db, tmp_path: Path, artifacts_dir: Path, reports_dir: Path
) -> None:
    org = _seed(db, tmp_path, artifacts_dir)
    ctx = build_context(db, org=org, period=parse_period("2026-04"), kind="fit-bundle")
    out = dx.write_report(ctx, reports_dir / "r.docx")

    doc = Document(str(out))
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert any("FIT-001" in h for h in headings)
    assert any("Impact" in h for h in headings)
    assert any("Remediation" in h for h in headings)
    assert any("Validation" in h for h in headings)


def test_write_single_ticket_docx(
    db, tmp_path: Path, artifacts_dir: Path, reports_dir: Path
) -> None:
    org = _seed(db, tmp_path, artifacts_dir)
    ctx = build_context(db, org=org, period=parse_period("2026-04"), kind="fit-bundle")
    assert ctx.tickets
    out = dx.write_ticket(ctx.tickets[0], reports_dir / "one.docx")
    assert out.exists()
